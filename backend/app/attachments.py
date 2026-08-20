"""Attachments: signed-URL upload to GCS, text extraction, video transcription.

Flow:
1. Frontend POSTs /attachments/init with {filename, content_type, size}.
   Backend validates against allowlist and size caps, generates a V4
   signed PUT URL pointing at GCS, persists a placeholder Firestore doc.
2. Frontend PUTs the file directly to GCS using the signed URL.
3. Frontend POSTs /attachments/{id}/complete. Backend verifies the object
   landed, extracts text (DOCX/TXT/CSV/MD) or transcribes (video) inline,
   and marks the attachment ready.
4. Chat references the attachment by id; chat.py rebuilds the OpenAI
   content parts at send time and on every history load.
"""
import base64
import io
import uuid
from datetime import timedelta
from typing import Annotated

from fastapi import APIRouter, Depends, HTTPException, status
from google.auth import default as google_default_credentials
from google.auth.transport.requests import Request as GoogleAuthRequest
from pydantic import BaseModel

from .auth import AuthedUser, require_user
from .config import get_settings
from .storage import archive_bucket, db, now

router = APIRouter()


# ---------------------- Allowlists & caps ----------------------

DOC_MIME = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "md",
    "text/x-markdown": "md",
    "text/csv": "csv",
}
IMAGE_MIME = {
    "image/png": "png",
    "image/jpeg": "jpg",
    "image/webp": "webp",
    "image/gif": "gif",
}
VIDEO_MIME = {
    "video/mp4": "mp4",
    "video/quicktime": "mov",
    "video/webm": "webm",
}

DOC_MAX = 10 * 1024 * 1024       # 10 MB
IMAGE_MAX = 5 * 1024 * 1024      # 5 MB (Anthropic image limit)
VIDEO_MAX = 100 * 1024 * 1024    # 100 MB

# ~4 chars per token, so this is roughly a 10k-token ceiling per document.
# The old 200k (~50k tokens) meant a single report could outweigh an entire
# conversation and, on its own, exceed a provider's per-minute token limit.
MAX_EXTRACTED_CHARS = 40_000


def _classify(content_type: str) -> tuple[str, int] | None:
    """Return (category, max_size_bytes) for an allowed mime; None otherwise."""
    if content_type in DOC_MIME:
        return ("document", DOC_MAX)
    if content_type in IMAGE_MIME:
        return ("image", IMAGE_MAX)
    if content_type in VIDEO_MIME:
        return ("video", VIDEO_MAX)
    return None


def _attachment_doc(uid: str, att_id: str):
    return (
        db()
        .collection("attachments")
        .document(uid)
        .collection("items")
        .document(att_id)
    )


def _refreshed_credentials():
    creds, _ = google_default_credentials()
    if not creds.valid:
        creds.refresh(GoogleAuthRequest())
    return creds


def _sign_url(blob, *, method: str, content_type: str | None = None,
              response_disposition: str | None = None, minutes: int = 15) -> str:
    """V4 signed URL using IAM SignBlob (no key file needed)."""
    creds = _refreshed_credentials()
    kwargs = dict(
        version="v4",
        expiration=timedelta(minutes=minutes),
        method=method,
        service_account_email=getattr(creds, "service_account_email", None),
        access_token=getattr(creds, "token", None),
    )
    if content_type is not None:
        kwargs["content_type"] = content_type
    if response_disposition is not None:
        kwargs["response_disposition"] = response_disposition
    return blob.generate_signed_url(**kwargs)


# ---------------------- Text extraction & transcription ----------------------

def _truncate(text: str) -> str:
    """Cap extracted text, telling the model when it is only seeing part."""
    if len(text) <= MAX_EXTRACTED_CHARS:
        return text
    return (
        text[:MAX_EXTRACTED_CHARS]
        + "\n\n[Document truncated here — this is the first "
        + f"{MAX_EXTRACTED_CHARS:,} characters only. Ask the user to re-attach "
        + "the relevant section if you need more.]"
    )


def _extract_pdf(raw: bytes) -> str:
    """Pull the text layer out of a PDF.

    Extracted at upload so later turns can reference a document by its text
    instead of re-sending the whole file. Scanned PDFs have no text layer and
    return "" — `build_content_parts` detects that and keeps sending those
    natively, since text is not an option for them.
    """
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(raw))
    pages: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            continue  # one unreadable page shouldn't lose the rest
        if text.strip():
            pages.append(text.strip())
    return _truncate("\n\n".join(pages))


def _extract_docx(raw: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text for cell in row.cells if cell.text]
            if row_cells:
                parts.append(" | ".join(row_cells))
    return _truncate("\n".join(parts))


def _extract_text_like(raw: bytes) -> str:
    try:
        return _truncate(raw.decode("utf-8", errors="replace"))
    except Exception:
        return _truncate(raw.decode("latin-1", errors="replace"))


def _transcribe_video(gcs_uri: str) -> str:
    """Transcribe a video's audio using Google Video Intelligence. Blocking.

    Returns a single concatenated transcript, or a sentinel string when no
    speech is detected. Pass-through of any client errors lets the /complete
    handler mark the attachment as errored.
    """
    from google.cloud import videointelligence_v1 as vi

    client = vi.VideoIntelligenceServiceClient()
    speech_cfg = vi.SpeechTranscriptionConfig(
        language_code="en-GB",
        enable_automatic_punctuation=True,
    )
    operation = client.annotate_video(
        request={
            "input_uri": gcs_uri,
            "features": [vi.Feature.SPEECH_TRANSCRIPTION],
            "video_context": vi.VideoContext(speech_transcription_config=speech_cfg),
        }
    )
    result = operation.result(timeout=900)  # 15 min ceiling

    pieces: list[str] = []
    for annotation in result.annotation_results:
        for transcription in annotation.speech_transcriptions:
            for alt in transcription.alternatives:
                if alt.transcript:
                    pieces.append(alt.transcript.strip())
                    break
    out = "\n".join(p for p in pieces if p).strip()
    return out or "[No speech detected in this video.]"


# ---------------------- Content-block construction ----------------------

def build_content_parts(
    text: str, attachments_full: list[dict], *, native: bool = True
) -> list[dict]:
    """Build the OpenAI Responses `input` content parts for one user message.

    `attachments_full` are full Firestore attachment records (with gcs_path,
    extracted_text, transcript). Video uses the stored transcript only.

    `native` controls PDF fidelity and is the lever that keeps long
    conversations affordable. A PDF sent natively is billed as its text *plus*
    an image per page — tens of thousands of tokens — and the old code paid
    that on every single turn for the life of the thread. With `native=False`
    the stored text layer is sent instead, which is a fraction of the size and
    still lets the model quote and reason about the document. Callers give
    full fidelity to the newest message only.

    Scanned PDFs have no text layer, so they stay native regardless — for
    those, downgrading would mean sending nothing at all.
    """
    bucket = archive_bucket()
    parts: list[dict] = []
    for att in attachments_full:
        ct = att.get("content_type", "")
        filename = att.get("filename") or "attachment"
        if ct == "application/pdf":
            pdf_text = att.get("extracted_text")
            if not native and pdf_text:
                parts.append({
                    "type": "input_text",
                    "text": (
                        f"[Attached document: {filename} — text of a PDF sent "
                        f"earlier in this conversation]\n\n{pdf_text}\n\n"
                        f"[End of attached document.]"
                    ),
                })
                continue
            raw = bucket.blob(att["gcs_path"]).download_as_bytes()
            b64 = base64.standard_b64encode(raw).decode("ascii")
            parts.append({
                "type": "input_file",
                "filename": filename,
                "file_data": f"data:application/pdf;base64,{b64}",
            })
        elif ct.startswith("image/"):
            raw = bucket.blob(att["gcs_path"]).download_as_bytes()
            b64 = base64.standard_b64encode(raw).decode("ascii")
            parts.append({
                "type": "input_image",
                "image_url": f"data:{ct};base64,{b64}",
                "detail": "auto",
            })
        elif att.get("extracted_text"):
            parts.append({
                "type": "input_text",
                "text": (
                    f"[Attached document: {filename}]\n\n"
                    f"{att['extracted_text']}\n\n"
                    f"[End of attached document.]"
                ),
            })
        elif att.get("transcript"):
            parts.append({
                "type": "input_text",
                "text": (
                    f"[Attached video: {filename}. Audio transcript follows. "
                    f"This is a transcript only — no biometric or emotion inference.]\n\n"
                    f"{att['transcript']}\n\n"
                    f"[End of transcript.]"
                ),
            })
    if text:
        parts.append({"type": "input_text", "text": text})
    return parts or [{"type": "input_text", "text": ""}]


def load_attachments_full(uid: str, attachment_ids: list[str]) -> list[dict]:
    """Load full attachment records for a user. Silently skips missing ids."""
    out: list[dict] = []
    for aid in attachment_ids or []:
        snap = _attachment_doc(uid, aid).get()
        if snap.exists:
            out.append(snap.to_dict())
    return out


def lightweight_attachment(att: dict) -> dict:
    """The subset stored on a message doc for rendering in the chat UI.

    `extracted_chars` also lets the history budgeter size a message without
    re-reading the attachment record from Firestore or GCS. Messages written
    before this field existed simply fall back to a size-based estimate.
    """
    return {
        "id": att.get("id"),
        "filename": att.get("filename"),
        "content_type": att.get("content_type"),
        "category": att.get("category"),
        "size": att.get("actual_size") or att.get("size"),
        "extracted_chars": len(att.get("extracted_text") or "") or None,
    }


def flag_scan_corpus(text: str, attachments_full: list[dict]) -> str:
    """Concatenate user text with extracted/transcribed text for the flag scan."""
    parts = [text or ""]
    for att in attachments_full:
        if att.get("extracted_text"):
            parts.append(att["extracted_text"])
        if att.get("transcript"):
            parts.append(att["transcript"])
    return "\n".join(parts)


# ---------------------- API surface ----------------------

class AttachmentInitBody(BaseModel):
    filename: str
    content_type: str
    size: int


def _serialize(att: dict) -> dict:
    return {
        "id": att.get("id"),
        "filename": att.get("filename"),
        "content_type": att.get("content_type"),
        "category": att.get("category"),
        "size": att.get("actual_size") or att.get("size"),
        "status": att.get("status"),
        "has_text": bool(att.get("extracted_text")),
        "has_transcript": bool(att.get("transcript")),
        "error": att.get("error"),
        "created_at": att["created_at"].isoformat() if att.get("created_at") else None,
        "ready_at": att["ready_at"].isoformat() if att.get("ready_at") else None,
    }


@router.post("/init")
async def init_upload(
    body: AttachmentInitBody,
    user: Annotated[AuthedUser, Depends(require_user)],
):
    cls = _classify(body.content_type)
    if cls is None:
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            f"Unsupported content type: {body.content_type}",
        )
    category, max_size = cls
    if body.size <= 0 or body.size > max_size:
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            f"File too large for {category} (max {max_size // (1024 * 1024)} MB)",
        )
    if len(body.filename) > 255 or not body.filename.strip():
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Invalid filename")

    att_id = "att_" + uuid.uuid4().hex[:20]
    safe_name = body.filename.replace("/", "_").replace("\\", "_")
    gcs_path = f"attachments/{user.uid}/{att_id}/{safe_name}"

    blob = archive_bucket().blob(gcs_path)
    upload_url = _sign_url(blob, method="PUT", content_type=body.content_type, minutes=15)

    _attachment_doc(user.uid, att_id).set({
        "id": att_id,
        "uid": user.uid,
        "filename": safe_name,
        "content_type": body.content_type,
        "size": body.size,
        "category": category,
        "gcs_path": gcs_path,
        "status": "uploading",
        "created_at": now(),
    })

    return {
        "id": att_id,
        "upload_url": upload_url,
        "required_headers": {"Content-Type": body.content_type},
        "expires_in": 15 * 60,
    }


@router.post("/{att_id}/complete")
async def complete_upload(
    att_id: str,
    user: Annotated[AuthedUser, Depends(require_user)],
):
    ref = _attachment_doc(user.uid, att_id)
    snap = ref.get()
    if not snap.exists:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Attachment not found")
    data = snap.to_dict()
    if data.get("status") == "ready":
        return _serialize(data)

    blob = archive_bucket().blob(data["gcs_path"])
    if not blob.exists():
        ref.update({"status": "error", "error": "Upload did not land in storage"})
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            "Upload did not land in storage — retry",
        )

    blob.reload()
    actual_size = blob.size or 0
    cls = _classify(data["content_type"])
    if cls is None or actual_size > cls[1]:
        try:
            blob.delete()
        finally:
            ref.update({"status": "error", "error": "Size exceeded after upload"})
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "File exceeded size limit")

    ref.update({"status": "processing"})
    update: dict = {"status": "ready", "ready_at": now(), "actual_size": actual_size}

    try:
        category = data["category"]
        ct = data["content_type"]
        if category == "document" and ct == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            update["extracted_text"] = _extract_docx(blob.download_as_bytes())
        elif category == "document" and ct in (
            "text/plain", "text/markdown", "text/x-markdown", "text/csv",
        ):
            update["extracted_text"] = _extract_text_like(blob.download_as_bytes())
        elif category == "document" and ct == "application/pdf":
            # Extracted once here so later turns can cite the document by text
            # rather than re-uploading it. Empty for scanned PDFs, which then
            # keep being sent natively.
            update["extracted_text"] = _extract_pdf(blob.download_as_bytes())
        elif category == "video":
            bucket_name = get_settings().GCS_ARCHIVE_BUCKET
            update["transcript"] = _transcribe_video(f"gs://{bucket_name}/{data['gcs_path']}")
        # Images have nothing to pre-extract — always sent natively.
    except Exception as e:
        ref.update({"status": "error", "error": str(e)[:500]})
        raise HTTPException(
            status.HTTP_500_INTERNAL_SERVER_ERROR,
            f"Attachment processing failed: {e}",
        )

    ref.update(update)
    return _serialize({**data, **update})


@router.get("/{att_id}")
async def get_attachment(
    att_id: str,
    user: Annotated[AuthedUser, Depends(require_user)],
):
    snap = _attachment_doc(user.uid, att_id).get()
    if not snap.exists:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Attachment not found")
    return _serialize(snap.to_dict())


@router.get("/{att_id}/download")
async def download_attachment(
    att_id: str,
    user: Annotated[AuthedUser, Depends(require_user)],
):
    snap = _attachment_doc(user.uid, att_id).get()
    if not snap.exists:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Attachment not found")
    data = snap.to_dict()
    blob = archive_bucket().blob(data["gcs_path"])
    if not blob.exists():
        raise HTTPException(status.HTTP_404_NOT_FOUND, "File not found in storage")
    url = _sign_url(
        blob,
        method="GET",
        response_disposition=f'attachment; filename="{data.get("filename", "file")}"',
        minutes=15,
    )
    return {"url": url, "expires_in": 15 * 60}
